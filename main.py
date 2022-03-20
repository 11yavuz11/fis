import sys
from PyQt5.QtWidgets import *
from pencere import *
from tkinter import *
import win32api
import win32print
import time
import random
from docx import Document
from docx.shared import Inches, Pt

Uygulama = QApplication(sys.argv)
pencere = QMainWindow()
ui = Ui_Pencere()
ui.setupUi(pencere)
pencere.show()
uyari = Errors()

global cikis
global dosya


def HESAPLA():
    try:
        document = Document()
        fiyat = float(ui.ln_fiyat.text())
        litre = float(ui.ln_litre.text())
        plaka = ui.ln_plaka.text()
        fisno = random.randint(0, 5000)
        pompano = int(ui.ln_pno.text())
        hesapla = litre * fiyat
        hesapla = round(hesapla, 2)
        kdv = hesapla / 100 * 18
        kdv = round(kdv, 2)
        cikis = "Plaka= " + str(plaka.upper()) + "\nFiyat= " + str(hesapla) + " TL\nKDV Fiyat= " + str(kdv) + "TL\nLitre= " + str(litre) + "\nPompa No: " + str(pompano)
        ui.lbl_cikti.setText(cikis)
        zaman = time.strftime('%x    SAAT %X')
        p = document.add_heading(level=0)
        p.add_run("""
//Fiş bilgileri
{} 
FIŞ NO: {}
                {}
{} LT X{}
VPD DIESEL %18     *{}

KDV                         *{}
TOP                         *{}

NAKİT                      *{}

POMPA 00{}
İSTASYON
                    """.format(zaman, fisno,plaka.upper(), litre, fiyat,hesapla,kdv, hesapla,hesapla,str(pompano))).font.size = Pt(8)
        run = p.add_run("")
        run.add_picture('mf.png', width=Inches(0.8))
        document.save("C:\\fis.docx")
        document.add_page_break()


    except:
        uyari.hesapla_error()


def TEMIZLE():
    ui.ln_fiyat.clear()
    ui.ln_litre.clear()
    ui.ln_plaka.clear()
    ui.lbl_cikti.clear()
    ui.ln_pno.clear()


def YAZDIR():
    try:
        filename = "C:\\fis.docx"
        open(filename, "r", encoding="utf-8")
        win32api.ShellExecute(
            0,
            "printto",
            filename,
            '"%s"' % win32print.GetDefaultPrinter(),
            ".",
            0
        )
    except:
        uyari.yazdir_error()


ui.btn_hesapla.clicked.connect(HESAPLA)
ui.btn_temizle.clicked.connect(TEMIZLE)
ui.btn_yazdir.clicked.connect(YAZDIR)
sys.exit(Uygulama.exec_())



